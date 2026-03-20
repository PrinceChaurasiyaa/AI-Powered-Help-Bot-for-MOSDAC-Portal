"""
crawler/document_parser.py
─────────────────────────────────────────────────────────────────────────────
Downloads and extracts text from documents linked on MOSDAC pages.

Supported formats
─────────────────
  PDF   — via pdfplumber (layout-aware) + PyPDF2 fallback
  DOCX  — via python-docx  (headings, paragraphs, tables)
  XLSX  — via openpyxl     (all sheets, all cells)
  CSV   — native Python csv module
  Scanned PDFs — OCR via pytesseract (if ENABLE_OCR=True)

Each parsed document is stored in the `documents` table with its full
extracted text.  The local file is also saved to DOCUMENTS_DIR.
─────────────────────────────────────────────────────────────────────────────
"""

import csv
import io
import os
from pathlib import Path
from typing import Optional, Tuple

import requests

from config import (
    DEFAULT_HEADERS,
    DOCUMENTS_DIR,
    ENABLE_OCR,
    MAX_DOC_SIZE_MB,
    REQUEST_TIMEOUT,
    SUPPORTED_DOC_TYPES,
)
from storage.data_store import DataStore
from utils.helpers import clean_text, content_hash, file_size_mb, url_hash
from utils.logger import get_logger

log = get_logger(__name__)


class DocumentParser:
    """
    Downloads a document from a URL and extracts its text content.

    Usage:
        store  = DataStore()
        parser = DocumentParser(store)
        parser.download_and_parse(
            "https://www.mosdac.gov.in/docs/user_manual.pdf",
            source_page_url="https://www.mosdac.gov.in/help"
        )
    """

    def __init__(self, store: DataStore):
        self.store   = store
        self.session = requests.Session()
        self.session.headers.update(DEFAULT_HEADERS)

    # ── Public API ────────────────────────────────────────────

    def download_and_parse(self, url: str, source_page_url: str = "") -> bool:
        """
        Downloads the document at `url`, extracts text, saves to DB.
        Returns True on success.
        """
        log.info(f"[Doc] Downloading: {url}")

        # ── Download ─────────────────────────────────────────
        raw_bytes, filename, file_type = self._download(url)
        if raw_bytes is None:
            self.store.mark_url_failed(url, "download failed")
            return False

        # ── Size check ───────────────────────────────────────
        size_kb = len(raw_bytes) / 1024
        size_mb = size_kb / 1024
        if size_mb > MAX_DOC_SIZE_MB:
            log.warning(f"[Doc] File too large ({size_mb:.1f} MB), skipping: {url}")
            self.store.mark_url_skipped(url, f"file too large: {size_mb:.1f} MB")
            return False

        # ── Save to disk ─────────────────────────────────────
        local_path = self._save_to_disk(raw_bytes, filename)

        # ── Extract text ─────────────────────────────────────
        extracted_text, page_count, extraction_ok = self._extract_text(
            raw_bytes, file_type, local_path
        )

        # ── Persist to DB ─────────────────────────────────────
        c_hash = content_hash(extracted_text)
        uid    = url_hash(url)

        self.store.save_document({
            "url":             url,
            "url_hash":        uid,
            "filename":        filename,
            "file_type":       file_type,
            "local_path":      str(local_path),
            "extracted_text":  extracted_text,
            "content_hash":    c_hash,
            "page_count":      page_count,
            "file_size_kb":    round(size_kb, 2),
            "source_page_url": source_page_url,
            "extraction_ok":   extraction_ok,
        })

        log.info(
            f"[Doc] ✓ {filename} | type={file_type} | "
            f"pages={page_count} | chars={len(extracted_text)}"
        )
        return True

    # ── Download ─────────────────────────────────────────────

    def _download(
        self, url: str
    ) -> Tuple[Optional[bytes], str, str]:
        """
        Returns (bytes, filename, file_type) or (None, '', '') on failure.
        """
        try:
            resp = self.session.get(
                url, timeout=REQUEST_TIMEOUT * 2,
                stream=True, allow_redirects=True
            )
            resp.raise_for_status()

            # Determine file extension from URL or Content-Type header
            from urllib.parse import urlparse
            url_path   = urlparse(url).path.lower()
            file_type  = self._detect_file_type(url_path, resp.headers)
            if not file_type:
                log.warning(f"[Doc] Unsupported file type for: {url}")
                return None, "", ""

            filename = Path(urlparse(url).path).name or f"document{file_type}"

            # Read in chunks to avoid memory issues for large files
            chunks = []
            for chunk in resp.iter_content(chunk_size=1024 * 256):  # 256 KB chunks
                chunks.append(chunk)
            raw_bytes = b"".join(chunks)

            return raw_bytes, filename, file_type

        except requests.exceptions.RequestException as exc:
            log.error(f"[Doc] Download error for {url}: {exc}")
            return None, "", ""

    def _detect_file_type(self, path: str, headers: dict) -> Optional[str]:
        """Detect file extension from URL path or Content-Type header."""
        for ext in SUPPORTED_DOC_TYPES:
            if path.endswith(ext):
                return ext

        ct = headers.get("Content-Type", "").lower()
        type_map = {
            "application/pdf":                                        ".pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml": ".docx",
            "application/msword":                                     ".doc",
            "application/vnd.openxmlformats-officedocument.spreadsheetml": ".xlsx",
            "application/vnd.ms-excel":                               ".xls",
            "text/csv":                                               ".csv",
        }
        for mime, ext in type_map.items():
            if mime in ct:
                return ext

        return None

    # ── Save to disk ─────────────────────────────────────────

    def _save_to_disk(self, data: bytes, filename: str) -> Path:
        """Write raw bytes to DOCUMENTS_DIR. Returns the saved path."""
        DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
        # Sanitise filename
        safe = "".join(c for c in filename if c.isalnum() or c in "._- ")[:200]
        dest = DOCUMENTS_DIR / safe

        # Avoid overwriting by appending a counter
        counter = 1
        while dest.exists():
            stem = Path(safe).stem
            ext  = Path(safe).suffix
            dest = DOCUMENTS_DIR / f"{stem}_{counter}{ext}"
            counter += 1

        dest.write_bytes(data)
        return dest

    # ── Text Extraction ──────────────────────────────────────

    def _extract_text(
        self, data: bytes, file_type: str, local_path: Path
    ) -> Tuple[str, int, bool]:
        """
        Dispatch to the correct parser based on file_type.
        Returns (extracted_text, page_count, success).
        """
        try:
            if file_type == ".pdf":
                return self._parse_pdf(data)
            elif file_type in (".docx", ".doc"):
                return self._parse_docx(local_path)
            elif file_type in (".xlsx", ".xls"):
                return self._parse_xlsx(local_path)
            elif file_type == ".csv":
                return self._parse_csv(data)
            else:
                return "", 0, False
        except Exception as exc:
            log.error(f"[Doc] Extraction failed ({file_type}): {exc}")
            return "", 0, False

    # ── PDF Parser ───────────────────────────────────────────

    def _parse_pdf(self, data: bytes) -> Tuple[str, int, bool]:
        """
        Strategy:
          1. Try pdfplumber (best for text-layer PDFs)
          2. Fall back to PyPDF2
          3. If both fail and ENABLE_OCR=True → pytesseract OCR
        """
        # ── pdfplumber ───────────────────────────────────────
        try:
            import pdfplumber
            pages_text = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)

            full_text = "\n\n".join(pages_text)
            if full_text.strip():
                log.debug(f"[PDF] pdfplumber: {page_count} pages, {len(full_text)} chars")
                return clean_text(full_text), page_count, True
        except Exception as e:
            log.debug(f"[PDF] pdfplumber failed: {e}")

        # ── PyPDF2 fallback ──────────────────────────────────
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            page_count = len(reader.pages)
            pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)

            full_text = "\n\n".join(pages_text)
            if full_text.strip():
                log.debug(f"[PDF] PyPDF2: {page_count} pages, {len(full_text)} chars")
                return clean_text(full_text), page_count, True
        except Exception as e:
            log.debug(f"[PDF] PyPDF2 failed: {e}")

        # ── OCR fallback (scanned PDFs) ───────────────────────
        if ENABLE_OCR:
            return self._ocr_pdf(data)

        log.warning("[PDF] Could not extract text (may be a scanned PDF). Enable OCR.")
        return "", 0, False

    def _ocr_pdf(self, data: bytes) -> Tuple[str, int, bool]:
        """
        Convert PDF pages to images and run Tesseract OCR.
        Requires:  pip install pdf2image pytesseract
        System dep: poppler, tesseract
        """
        try:
            from pdf2image import convert_from_bytes
            import pytesseract

            images = convert_from_bytes(data, dpi=200)
            pages_text = []
            for img in images:
                text = pytesseract.image_to_string(img, lang="eng")
                pages_text.append(text)

            full_text = "\n\n".join(pages_text)
            log.info(f"[PDF][OCR] Extracted {len(pages_text)} pages via OCR")
            return clean_text(full_text), len(images), True
        except Exception as exc:
            log.error(f"[PDF][OCR] Failed: {exc}")
            return "", 0, False

    # ── DOCX Parser ──────────────────────────────────────────

    def _parse_docx(self, path: Path) -> Tuple[str, int, bool]:
        """
        Extract text from Word documents.
        Preserves heading structure and table content.
        """
        try:
            from docx import Document
            doc = Document(str(path))

            parts = []

            # Paragraphs (includes headings)
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                # Mark headings
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    parts.append(f"{'#' * int(level) if level.isdigit() else '#'} {text}")
                else:
                    parts.append(text)

            # Tables within the document
            for table in doc.tables:
                rows_text = []
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    rows_text.append(" | ".join(cells))
                if rows_text:
                    parts.append("\n".join(rows_text))

            full_text = "\n\n".join(parts)
            log.debug(f"[DOCX] Extracted {len(full_text)} chars")
            return clean_text(full_text), len(doc.paragraphs), True
        except Exception as exc:
            log.error(f"[DOCX] Parse error: {exc}")
            return "", 0, False

    # ── XLSX Parser ──────────────────────────────────────────

    def _parse_xlsx(self, path: Path) -> Tuple[str, int, bool]:
        """
        Extract all sheets from Excel files.
        Returns a structured text representation.
        """
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)

            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"=== Sheet: {sheet_name} ===")
                for row in ws.iter_rows(values_only=True):
                    # Filter empty rows
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        parts.append(" | ".join(cells))

            full_text = "\n".join(parts)
            log.debug(f"[XLSX] Extracted {len(wb.sheetnames)} sheets, {len(full_text)} chars")
            return clean_text(full_text), len(wb.sheetnames), True
        except Exception as exc:
            log.error(f"[XLSX] Parse error: {exc}")
            return "", 0, False

    # ── CSV Parser ───────────────────────────────────────────

    def _parse_csv(self, data: bytes) -> Tuple[str, int, bool]:
        """Parse CSV files into pipe-delimited text."""
        try:
            # Detect encoding
            import chardet
            detected = chardet.detect(data)
            encoding = detected.get("encoding") or "utf-8"

            text_io = io.StringIO(data.decode(encoding, errors="replace"))
            reader  = csv.reader(text_io)
            rows    = list(reader)
            lines   = [" | ".join(row) for row in rows if any(c.strip() for c in row)]

            full_text = "\n".join(lines)
            log.debug(f"[CSV] Extracted {len(rows)} rows, {len(full_text)} chars")
            return clean_text(full_text), len(rows), True
        except Exception as exc:
            log.error(f"[CSV] Parse error: {exc}")
            return "", 0, False
